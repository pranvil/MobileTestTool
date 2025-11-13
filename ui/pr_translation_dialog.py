#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PR翻译对话框
"""

import os
import urllib.parse
import urllib.request
import json
import datetime
import html
from PyQt5.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QLabel, 
                             QLineEdit, QTextEdit, QPushButton, QMessageBox,
                             QScrollArea, QWidget, QFrame, QComboBox, QApplication)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt5.QtGui import QTextCursor
import subprocess
from ui.widgets.shadow_utils import add_card_shadow
from core.debug_logger import logger


class PlainTextEdit(QTextEdit):
    """只允许粘贴纯文本的QTextEdit"""
    
    def insertFromMimeData(self, source):
        """重写粘贴方法，只粘贴纯文本"""
        if source.hasText():
            # 只获取纯文本内容
            plain_text = source.text()
            # 插入纯文本
            self.textCursor().insertText(plain_text)
        # 忽略其他格式（HTML、图片等）


try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TranslationWorker(QThread):
    """翻译工作线程"""
    finished = pyqtSignal(str)  # 翻译完成，传递文件路径
    error = pyqtSignal(str)  # 翻译错误
    progress = pyqtSignal(str)  # 进度更新
    
    def __init__(self, data_dict, email=None, translate_direction='zh_en', parent=None):
        super().__init__(parent)
        self.data_dict = data_dict
        self.email = email
        self.translate_direction = translate_direction  # 'zh_en' 或 'en_zh'
        self.lang_manager = parent.lang_manager if parent and hasattr(parent, 'lang_manager') else None
    
    def tr(self, text):
        """安全地获取翻译文本"""
        return self.lang_manager.tr(text) if self.lang_manager else text
    
    def run(self):
        """执行翻译"""
        try:
            # 定义不需要翻译的字段
            no_translate_fields = ['log', 'associate_specification', 'test_plan_reference', 
                                  'tools_and_platforms', 'user_impact', 'reproducing_rate']
            
            # 根据翻译方向处理数据
            # Word文档始终是中文在前，英文在后
            if self.translate_direction == 'zh_en':
                # 中文翻译英文：输入的是中文，翻译成英文
                source_dict = self.data_dict  # 中文内容
                translations = {}
                for key, source_text in source_dict.items():
                    if source_text and source_text.strip():
                        if key in no_translate_fields:
                            translations[key] = ""
                            self.progress.emit(f"{self.tr('跳过翻译')}: {key}...")
                        else:
                            self.progress.emit(f"{self.tr('正在翻译')}: {key}...")
                            translated = self.translate_text(source_text)
                            translations[key] = translated
                    else:
                        translations[key] = ""
                chinese_dict = source_dict
                english_dict = translations
            else:
                # 英文翻译中文：输入的是英文，翻译成中文
                source_dict = self.data_dict  # 英文内容
                translations = {}
                for key, source_text in source_dict.items():
                    if source_text and source_text.strip():
                        if key in no_translate_fields:
                            translations[key] = ""
                            self.progress.emit(f"{self.tr('跳过翻译')}: {key}...")
                        else:
                            self.progress.emit(f"{self.tr('正在翻译')}: {key}...")
                            translated = self.translate_text(source_text)
                            translations[key] = translated
                    else:
                        translations[key] = ""
                # Word文档格式：中文在前，英文在后
                # 所以英文输入时，翻译后的中文作为"中文"，原始英文作为"英文"
                chinese_dict = translations  # 翻译后的中文
                english_dict = source_dict  # 原始英文
            
            # 生成Word文档
            self.progress.emit(self.tr("正在生成Word文档..."))
            doc_path = self.generate_word_document(chinese_dict, english_dict)
            
            self.finished.emit(doc_path)
        except Exception as e:
            logger.exception(f"{self.tr('翻译失败')}: {str(e)}")
            self.error.emit(str(e))
    
    def translate_text(self, text):
        """使用MyMemory API翻译文本"""
        # MyMemory API限制：单次查询最大500字符
        # 为了安全起见，使用450字符作为分割点
        MAX_CHARS = 450
        
        # 如果文本为空，直接返回
        if not text or not text.strip():
            return text
        
        # 如果文本长度在限制内，直接翻译
        if len(text) <= MAX_CHARS:
            return self._translate_single_text(text)
        
        # 文本过长，需要分段翻译
        # 使用更智能的分割策略：按段落分割，但如果段落太长则按句子分割
        # 先尝试按双换行符分割段落
        parts = []
        remaining_text = text
        
        while remaining_text:
            # 如果剩余文本在限制内，直接翻译
            if len(remaining_text) <= MAX_CHARS:
                translated = self._translate_single_text(remaining_text)
                parts.append({
                    'text': translated,
                    'original_end': len(remaining_text)
                })
                break
            
            # 尝试找到最佳分割点（在MAX_CHARS附近）
            # 优先在段落边界（双换行符）分割
            split_pos = MAX_CHARS
            best_split = None
            
            # 查找最近的段落边界（向前查找）
            para_boundary = remaining_text.rfind('\n\n', 0, MAX_CHARS)
            if para_boundary > MAX_CHARS * 0.5:  # 如果段落边界在合理范围内
                best_split = para_boundary + 2  # 包含双换行符
            else:
                # 如果没有找到段落边界，尝试在单换行符处分割
                line_boundary = remaining_text.rfind('\n', 0, MAX_CHARS)
                if line_boundary > MAX_CHARS * 0.5:
                    best_split = line_boundary + 1
                else:
                    # 如果都没有，尝试在句子边界分割
                    sentence_boundary = self._find_sentence_boundary(remaining_text, MAX_CHARS)
                    if sentence_boundary > MAX_CHARS * 0.5:
                        best_split = sentence_boundary
                    else:
                        # 最后选择在空格处分割
                        space_boundary = remaining_text.rfind(' ', 0, MAX_CHARS)
                        if space_boundary > MAX_CHARS * 0.5:
                            best_split = space_boundary + 1
                        else:
                            # 强制在MAX_CHARS处分割
                            best_split = MAX_CHARS
            
            # 分割并翻译第一部分（保留末尾空白，但翻译时去掉）
            chunk_with_whitespace = remaining_text[:best_split]
            chunk = chunk_with_whitespace.rstrip()
            
            if chunk:
                translated_chunk = self._translate_single_text(chunk)
                
                # 提取chunk后面的空白字符作为分隔符
                trailing_whitespace = chunk_with_whitespace[len(chunk):]
                
                parts.append({
                    'text': translated_chunk,
                    'separator': trailing_whitespace if trailing_whitespace else ''
                })
            
            # 更新剩余文本
            remaining_text = remaining_text[best_split:]
        
        # 合并翻译结果，保留原始格式
        result_parts = []
        for i, part in enumerate(parts):
            result_parts.append(part['text'])
            # 添加分隔符（除了最后一部分）
            if i < len(parts) - 1 and 'separator' in part:
                result_parts.append(part['separator'])
        
        return ''.join(result_parts)
    
    def _find_sentence_boundary(self, text, max_pos):
        """查找句子边界（句号、问号、感叹号后跟空格）"""
        import re
        # 查找句号、问号、感叹号后跟空格或换行的位置
        pattern = r'[.!?]\s+'
        matches = list(re.finditer(pattern, text[:max_pos]))
        if matches:
            # 返回最后一个匹配的结束位置
            return matches[-1].end()
        return max_pos
    
    def _split_into_sentences(self, text):
        """将文本按句子分割"""
        import re
        # 按句号、问号、感叹号分割，保留分隔符
        sentences = re.split(r'([.!?]\s+)', text)
        result = []
        i = 0
        while i < len(sentences):
            if i + 1 < len(sentences):
                result.append(sentences[i] + sentences[i + 1])
                i += 2
            else:
                if sentences[i].strip():
                    result.append(sentences[i])
                i += 1
        return result
    
    def _translate_single_text(self, text):
        """翻译单个文本片段（不超过500字符）"""
        try:
            # 根据翻译方向设置语言对
            if self.translate_direction == 'zh_en':
                langpair = 'zh|en'  # 中文翻译成英文
            else:
                langpair = 'en|zh'  # 英文翻译成中文
            
            # 构建API URL
            base_url = "https://api.mymemory.translated.net/get"
            params = {
                'q': text,
                'langpair': langpair
            }
            
            # 如果提供了邮箱，添加到参数中
            if self.email and self.email.strip():
                params['de'] = self.email.strip()
            
            # 使用urlencode确保特殊字符（如->）正确编码
            # encoding='utf-8'确保中文字符正确编码
            # quote_via=urllib.parse.quote_plus 使用quote_plus处理空格和特殊字符
            url = f"{base_url}?{urllib.parse.urlencode(params, encoding='utf-8', quote_via=urllib.parse.quote_plus)}"
            
            # 发送请求
            request = urllib.request.Request(url)
            request.add_header('User-Agent', 'Mozilla/5.0')
            request.add_header('Accept-Charset', 'UTF-8')
            
            with urllib.request.urlopen(request, timeout=30) as response:
                # 确保使用UTF-8解码响应
                response_data = response.read()
                decoded_data = response_data.decode('utf-8', errors='replace')
                data = json.loads(decoded_data)
                
                if data.get('responseStatus') == 200:
                    translated_text = data.get('responseData', {}).get('translatedText', '')
                    # 解码HTML实体（如 &lt; -> <, &gt; -> >, &amp; -> &）
                    translated_text = html.unescape(translated_text)
                    return translated_text
                else:
                    error_msg = data.get('responseDetails', '')
                    logger.warning(f"MyMemory API错误: {data.get('responseStatus')} - {error_msg}")
                    # 如果是长度超限错误，尝试进一步分割
                    if '500' in error_msg or 'LENGTH' in error_msg.upper():
                        logger.warning(f"文本仍然过长，尝试进一步分割: {len(text)} 字符")
                        # 递归分割
                        mid = len(text) // 2
                        # 尝试在空格处分割
                        split_pos = text.rfind(' ', 0, mid)
                        if split_pos == -1:
                            split_pos = mid
                        part1 = self._translate_single_text(text[:split_pos])
                        part2 = self._translate_single_text(text[split_pos:].lstrip())
                        return part1 + " " + part2
                    return text  # 翻译失败时返回原文
        except Exception as e:
            logger.exception(f"翻译API调用失败: {str(e)}")
            return text  # 翻译失败时返回原文
    
    def generate_word_document(self, chinese_dict, english_dict):
        """生成Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，请运行: pip install python-docx")
        
        # 创建文档
        doc = Document()
        
        # 定义标题映射
        title_map = {
            'summary': '概要/Summary',
            'defect_description': '缺陷描述/DEFECT DESCRIPTION',
            'preconditions': '前置条件/Preconditions',
            'steps': '步骤/Steps',
            'actual_result': '实际结果/Actual result',
            'expected_result': '预期结果/Expected result',
            'log': '日志如下/Log as below',
            'associate_specification': '关联规范/ASSOCIATE SPECIFICATION',
            'test_plan_reference': '测试计划参考/TEST PLAN REFERENCE',
            'tools_and_platforms': '使用的工具和平台/TOOLS AND PLATFORMS USED',
            'user_impact': '用户影响/USER IMPACT',
            'reproducing_rate': '复现率/REPRODUCING RATE',
            'reference_mobile': '对于场测问题（FT PR），请列出参考机的表现/For FT PR,Please list reference mobile\'s behavior'
        }
        
        # 遍历所有字段
        for key, title in title_map.items():
            chinese_text = chinese_dict.get(key, '').strip()
            english_text = english_dict.get(key, '').strip()
            
            # 添加标题（加粗，后面加冒号）
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{title}:")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            # 添加内容
            if chinese_text:
                # 添加中文内容（字号比标题小两个字号，即10pt）
                chinese_para = doc.add_paragraph(chinese_text)
                chinese_para.paragraph_format.space_after = Pt(6)
                for run in chinese_para.runs:
                    run.font.size = Pt(10)
                
                # 添加英文翻译
                if english_text:
                    if key == 'summary':
                        # 概要：中文后紧跟英文，不换行
                        chinese_run = chinese_para.runs[0]
                        chinese_run.add_text(f" / {english_text}")
                    else:
                        # 其他字段：换行显示英文（字号10pt）
                        english_para = doc.add_paragraph(english_text)
                        english_para.paragraph_format.space_after = Pt(12)
                        for run in english_para.runs:
                            run.font.size = Pt(10)
            
            # 添加段落间距（在标题和内容之后）
            doc.add_paragraph()
        
        # 获取存储路径（和log的存储路径逻辑一致）
        storage_path = self.get_storage_path()
        
        # 确保目录存在
        if not os.path.exists(storage_path):
            os.makedirs(storage_path)
        
        # 保存文档
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"PR_Translation_{timestamp}.docx"
        file_path = os.path.join(storage_path, filename)
        doc.save(file_path)
        
        return file_path
    
    def get_storage_path(self):
        """获取存储路径，优先使用用户配置的路径"""
        # 从父窗口（PRTranslationDialog）获取主窗口的工具配置
        parent_dialog = self.parent()
        if parent_dialog and hasattr(parent_dialog, 'parent'):
            main_window = parent_dialog.parent()
            if main_window and hasattr(main_window, 'tool_config') and main_window.tool_config:
                storage_path = main_window.tool_config.get("storage_path", "")
                if storage_path:
                    return storage_path
        
        # 使用默认路径
        current_date = datetime.datetime.now().strftime("%Y%m%d")
        return f"c:\\log\\{current_date}"


class PRTranslationDialog(QDialog):
    """PR翻译对话框"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        # 从父窗口获取语言管理器
        self.lang_manager = parent.lang_manager if parent and hasattr(parent, 'lang_manager') else None
        self.translation_worker = None
        # 配置文件路径
        self.config_file = os.path.expanduser("~/.netui/pr_translation_config.json")
        self.setup_ui()
        # 加载保存的邮箱
        self.load_saved_email()
    
    def tr(self, text):
        """安全地获取翻译文本"""
        return self.lang_manager.tr(text) if self.lang_manager else text
    
    def setup_ui(self):
        """设置UI"""
        self.setWindowTitle(self.tr("PR翻译"))
        self.setModal(True)
        self.resize(800, 700)
        
        # 设置窗口标志，允许最大化
        self.setWindowFlags(self.windowFlags() | Qt.WindowMaximizeButtonHint)
        
        # 主布局
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)
        
        # 顶部工具栏：翻译方向选择
        toolbar_layout = QHBoxLayout()
        toolbar_layout.setSpacing(10)
        
        # 翻译方向选择
        direction_label = QLabel(self.tr("翻译方向:"))
        self.direction_combo = QComboBox()
        self.direction_combo.addItem(self.tr("中文 → 英文"), 'zh_en')
        self.direction_combo.addItem(self.tr("英文 → 中文"), 'en_zh')
        self.direction_combo.setCurrentIndex(0)  # 默认中文翻译英文
        toolbar_layout.addWidget(direction_label)
        toolbar_layout.addWidget(self.direction_combo)
        
        toolbar_layout.addStretch()
        
        main_layout.addLayout(toolbar_layout)
        
        # 创建滚动区域
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        
        # 创建滚动内容容器
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(5, 5, 5, 5)
        scroll_layout.setSpacing(10)
        
        # 邮箱输入框
        email_container = QWidget()
        email_layout = QVBoxLayout(email_container)
        email_layout.setContentsMargins(0, 0, 0, 0)
        email_layout.setSpacing(4)
        
        email_title = QLabel(self.tr("邮箱（可选）"))
        email_title.setProperty("class", "section-title")
        email_layout.addWidget(email_title)
        
        email_card = QFrame()
        email_card.setObjectName("card")
        add_card_shadow(email_card)
        email_card_layout = QVBoxLayout(email_card)
        email_card_layout.setContentsMargins(10, 1, 10, 1)
        
        self.email_edit = QLineEdit()
        self.email_edit.setPlaceholderText(self.tr("输入邮箱地址（可选，用于提高翻译配额）"))
        # 监听邮箱输入变化，自动保存
        self.email_edit.textChanged.connect(self._on_email_changed)
        email_card_layout.addWidget(self.email_edit)
        
        email_layout.addWidget(email_card)
        scroll_layout.addWidget(email_container)
        
        # 定义字段配置：字段名 -> (标题, 是否多行)
        fields_config = {
            'summary': (self.tr('概要/Summary'), False),
            'defect_description': (self.tr('缺陷描述/DEFECT DESCRIPTION'), False),
            'preconditions': (self.tr('前置条件/Preconditions'), False),
            'steps': (self.tr('步骤/Steps'), True),
            'actual_result': (self.tr('实际结果/Actual result'), True),
            'expected_result': (self.tr('预期结果/Expected result'), True),
            'log': (self.tr('日志如下/Log as below'), False),
            'associate_specification': (self.tr('关联规范/ASSOCIATE SPECIFICATION'), False),
            'test_plan_reference': (self.tr('测试计划参考/TEST PLAN REFERENCE'), False),
            'tools_and_platforms': (self.tr('使用的工具和平台/TOOLS AND PLATFORMS USED'), False),
            'user_impact': (self.tr('用户影响/USER IMPACT'), False),
            'reproducing_rate': (self.tr('复现率/REPRODUCING RATE'), False),
            'reference_mobile': (self.tr('对于场测问题（FT PR），请列出参考机的表现/For FT PR,Please list reference mobile\'s behavior'), False)
        }
        
        # 存储输入框引用
        self.input_fields = {}
        
        # 创建每个字段的输入区域
        for field_key, (field_title, is_multiline) in fields_config.items():
            field_container = QWidget()
            field_layout = QVBoxLayout(field_container)
            field_layout.setContentsMargins(0, 0, 0, 0)
            field_layout.setSpacing(4)
            
            # 标题
            title_label = QLabel(field_title)
            title_label.setProperty("class", "section-title")
            field_layout.addWidget(title_label)
            
            # 卡片容器
            field_card = QFrame()
            field_card.setObjectName("card")
            add_card_shadow(field_card)
            field_card_layout = QVBoxLayout(field_card)
            field_card_layout.setContentsMargins(10, 1, 10, 1)
            
            # 输入框（所有输入框都使用PlainTextEdit，只允许粘贴纯文本）
            input_widget = PlainTextEdit()
            if is_multiline:
                # 多行输入框：默认高度100，但可以根据内容自动调整
                input_widget.setMinimumHeight(100)
                input_widget.setMaximumHeight(500)  # 设置最大高度，避免无限增长
                # 连接文本变化信号，自动调整高度
                input_widget.textChanged.connect(
                    lambda widget=input_widget: self._adjust_multiline_text_edit_height(widget)
                )
                # 初始设置高度为100像素
                input_widget.setFixedHeight(100)
            else:
                # 单行输入框：默认一行高度，但可以根据内容自动调整
                input_widget.setMinimumHeight(25)
                input_widget.setMaximumHeight(200)  # 设置最大高度，避免无限增长
                # 连接文本变化信号，自动调整高度
                # 使用lambda捕获input_widget的引用（通过默认参数）
                input_widget.textChanged.connect(
                    lambda widget=input_widget: self._adjust_text_edit_height(widget)
                )
                # 初始设置高度为一行
                input_widget.setFixedHeight(25)
            
            field_card_layout.addWidget(input_widget)
            field_layout.addWidget(field_card)
            scroll_layout.addWidget(field_container)
            
            # 保存引用
            self.input_fields[field_key] = input_widget
        
        # 设置滚动区域的内容
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area)
        
        # 底部按钮
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        self.translate_btn = QPushButton(self.tr("确认翻译"))
        self.translate_btn.clicked.connect(self.on_translate)
        button_layout.addWidget(self.translate_btn)
        
        self.cancel_btn = QPushButton(self.tr("取消"))
        self.cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(self.cancel_btn)
        
        main_layout.addLayout(button_layout)
    
    def on_translate(self):
        """开始翻译"""
        # 检查python-docx是否可用
        if not DOCX_AVAILABLE:
            QMessageBox.warning(
                self, 
                self.tr("缺少依赖"), 
                self.tr("需要安装python-docx库才能生成Word文档。\n请运行: pip install python-docx")
            )
            return
        
        # 收集所有输入数据（所有输入框都是QTextEdit）
        data_dict = {}
        for field_key, input_widget in self.input_fields.items():
            text = input_widget.toPlainText().strip()
            data_dict[field_key] = text
        
        # 检查是否有任何内容需要翻译
        has_content = any(data_dict.values())
        if not has_content:
            QMessageBox.warning(self, self.tr("输入错误"), self.tr("请至少输入一个字段的内容"))
            return
        
        # 获取邮箱并保存
        email = self.email_edit.text().strip() if self.email_edit.text().strip() else None
        if email:
            self.save_email(email)
        
        # 获取翻译方向
        translate_direction = self.direction_combo.currentData()
        
        # 禁用按钮
        self.translate_btn.setEnabled(False)
        self.cancel_btn.setEnabled(False)
        
        # 创建进度对话框
        self.progress_dialog = QMessageBox(self)
        self.progress_dialog.setWindowTitle(self.tr("翻译中"))
        self.progress_dialog.setText(self.tr("正在翻译，请稍候..."))
        self.progress_dialog.setStandardButtons(QMessageBox.NoButton)
        self.progress_dialog.show()
        
        # 创建工作线程
        self.translation_worker = TranslationWorker(data_dict, email, translate_direction, self)
        self.translation_worker.finished.connect(self.on_translation_finished)
        self.translation_worker.error.connect(self.on_translation_error)
        self.translation_worker.progress.connect(self.on_translation_progress)
        self.translation_worker.start()
    
    def on_translation_progress(self, message):
        """翻译进度更新"""
        if hasattr(self, 'progress_dialog'):
            self.progress_dialog.setText(message)
    
    def on_translation_finished(self, file_path):
        """翻译完成"""
        # 彻底关闭进度对话框，避免阻塞
        if hasattr(self, 'progress_dialog') and self.progress_dialog:
            self.progress_dialog.hide()  # 先隐藏
            self.progress_dialog.close()  # 关闭
            self.progress_dialog.deleteLater()  # 确保彻底销毁
            self.progress_dialog = None  # 清除引用
            # 强制处理事件，确保UI立即更新
            QApplication.processEvents()
        
        # 恢复按钮
        self.translate_btn.setEnabled(True)
        self.cancel_btn.setEnabled(True)
        
        # 使用QTimer延迟打开文件，确保UI更新完成后再打开
        # 这样可以避免阻塞主线程
        QTimer.singleShot(200, lambda: self._open_word_document(file_path))
    
    def _open_word_document(self, file_path):
        """打开Word文档（非阻塞方式）"""
        try:
            # 再次确保进度对话框已关闭
            if hasattr(self, 'progress_dialog') and self.progress_dialog:
                self.progress_dialog.hide()
                self.progress_dialog.close()
                self.progress_dialog.deleteLater()
                self.progress_dialog = None
                QApplication.processEvents()
            
            # 使用subprocess.Popen非阻塞方式打开文件
            if os.name == 'nt':  # Windows
                subprocess.Popen(
                    ['start', '', file_path],
                    shell=True,
                    creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
                )
            else:
                # 其他系统使用os.startfile或xdg-open
                os.startfile(file_path)
            
            # 再次处理事件，确保文件打开操作完成
            QApplication.processEvents()
            
            # 延迟显示成功消息，确保进度对话框完全消失
            QTimer.singleShot(100, lambda: self._show_success_message(file_path))
        except Exception as e:
            # 确保进度对话框已关闭后再显示错误消息
            if hasattr(self, 'progress_dialog') and self.progress_dialog:
                self.progress_dialog.hide()
                self.progress_dialog.close()
                self.progress_dialog.deleteLater()
                self.progress_dialog = None
                QApplication.processEvents()
            
            QMessageBox.critical(
                self, 
                self.tr("打开失败"), 
                self.tr(f"翻译完成，但打开文档失败: {str(e)}\n文件路径: {file_path}")
            )
    
    def _show_success_message(self, file_path):
        """显示成功消息"""
        # 最后一次确保进度对话框已关闭
        if hasattr(self, 'progress_dialog') and self.progress_dialog:
            self.progress_dialog.hide()
            self.progress_dialog.close()
            self.progress_dialog.deleteLater()
            self.progress_dialog = None
            QApplication.processEvents()
        
        QMessageBox.information(
            self, 
            self.tr("翻译完成"), 
            self.tr(f"翻译完成！Word文档已生成并打开。\n文件路径: {file_path}")
        )
        # 不自动关闭对话框，让用户可以继续使用
    
    def _adjust_text_edit_height(self, text_edit):
        """根据内容自动调整单行QTextEdit的高度"""
        # 获取文档内容高度
        doc_height = text_edit.document().size().height()
        
        # 添加一些边距（上下各4像素）
        margin = 8
        new_height = int(doc_height) + margin
        
        # 确保高度在最小和最大范围内
        min_height = 25
        max_height = 200
        new_height = max(min_height, min(new_height, max_height))
        
        # 设置新高度
        text_edit.setFixedHeight(new_height)
    
    def _adjust_multiline_text_edit_height(self, text_edit):
        """根据内容自动调整多行QTextEdit的高度"""
        # 获取文档内容高度
        doc_height = text_edit.document().size().height()
        
        # 添加一些边距（上下各4像素）
        margin = 8
        new_height = int(doc_height) + margin
        
        # 确保高度在最小和最大范围内
        min_height = 100  # 多行输入框最小高度
        max_height = 500  # 多行输入框最大高度
        new_height = max(min_height, min(new_height, max_height))
        
        # 设置新高度
        text_edit.setFixedHeight(new_height)
    
    def load_saved_email(self):
        """加载保存的邮箱"""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    saved_email = config.get('email', '')
                    if saved_email:
                        self.email_edit.setText(saved_email)
        except Exception as e:
            logger.exception(f"{self.tr('加载保存的邮箱失败')}: {str(e)}")
    
    def save_email(self, email):
        """保存邮箱到配置文件"""
        try:
            # 确保配置目录存在
            os.makedirs(os.path.dirname(self.config_file), exist_ok=True)
            
            # 读取现有配置或创建新配置
            config = {}
            if os.path.exists(self.config_file):
                try:
                    with open(self.config_file, 'r', encoding='utf-8') as f:
                        config = json.load(f)
                except Exception:
                    config = {}
            
            # 更新邮箱
            config['email'] = email
            
            # 保存配置
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.exception(f"{self.tr('保存邮箱失败')}: {str(e)}")
    
    def _on_email_changed(self, text):
        """邮箱输入变化时的处理（延迟保存，避免频繁写入）"""
        # 使用QTimer延迟保存，避免每次输入都保存
        if not hasattr(self, '_email_save_timer'):
            self._email_save_timer = QTimer()
            self._email_save_timer.setSingleShot(True)
            self._email_save_timer.timeout.connect(lambda: self._save_email_if_not_empty())
        
        # 重置定时器，延迟1秒后保存
        self._email_save_timer.stop()
        self._email_save_timer.start(1000)
    
    def _save_email_if_not_empty(self):
        """保存邮箱（如果非空）"""
        email = self.email_edit.text().strip()
        if email:
            self.save_email(email)
    
    def on_translation_error(self, error_msg):
        """翻译错误"""
        # 关闭进度对话框
        if hasattr(self, 'progress_dialog'):
            self.progress_dialog.close()
        
        # 恢复按钮
        self.translate_btn.setEnabled(True)
        self.cancel_btn.setEnabled(True)
        
        QMessageBox.critical(self, self.tr("翻译失败"), self.tr(f"翻译过程中发生错误: {error_msg}"))

